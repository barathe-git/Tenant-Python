from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Query, Form
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
import os
import uuid
from pathlib import Path
from typing import Optional
from datetime import date, datetime
from dateutil.relativedelta import relativedelta
from backend.models.database import get_db
from backend.models.models import Tenant, Owner, Building, Client
from backend.auth.auth import get_current_user
from dotenv import load_dotenv
import logging
from docx import Document
import re

load_dotenv()

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/files", tags=["files"])

UPLOAD_DIR = os.getenv("UPLOAD_DIR", "uploads")
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB

# Create upload directory if it doesn't exist
Path(UPLOAD_DIR).mkdir(parents=True, exist_ok=True)


@router.post("/upload")
def upload_file(
    file: UploadFile = File(...),
    tenant_id: Optional[int] = Query(None),
    file_type: str = Query("agreement", description="Type of file: 'agreement' or 'aadhar'"),
    db: Session = Depends(get_db)
):
    """Upload a PDF file"""
    # Log received parameters
    logger.info(f"Upload request - tenant_id: {tenant_id}, file_type: '{file_type}', filename: {file.filename}")

    # Validate file type
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are allowed")

    # Validate file_type parameter
    file_type = file_type.strip().lower()  # Normalize the file_type
    if file_type not in ["agreement", "aadhar"]:
        logger.error(f"Invalid file_type received: '{file_type}'")
        raise HTTPException(status_code=400, detail=f"file_type must be 'agreement' or 'aadhar', got '{file_type}'")

    # Read file content to check size
    file_content = file.file.read()
    file_size = len(file_content)

    if file_size > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail=f"File size exceeds maximum allowed size of {MAX_FILE_SIZE / 1024 / 1024}MB")

    # Generate meaningful filename based on tenant_id and file_type
    file_extension = Path(file.filename).suffix.lower()
    if tenant_id:
        # Use tenant_id and file_type for meaningful naming: tenant_1_agreement.pdf or tenant_1_aadhar.pdf
        unique_filename = f"tenant_{tenant_id}_{file_type}{file_extension}"
    else:
        # Fallback to UUID if no tenant_id provided
        unique_filename = f"{uuid.uuid4()}{file_extension}"
    file_path = os.path.join(UPLOAD_DIR, unique_filename)

    logger.info(f"Saving file to: {file_path}")

    # Save file
    with open(file_path, "wb") as f:
        f.write(file_content)

    # Update tenant record if tenant_id provided
    if tenant_id:
        tenant = db.query(Tenant).filter(Tenant.tenant_id == tenant_id).first()
        if tenant:
            logger.info(f"Updating tenant {tenant_id} with file_type '{file_type}'")
            if file_type == "agreement":
                # Delete old agreement file if exists and different
                if tenant.agreement_pdf_path and tenant.agreement_pdf_path != file_path and os.path.exists(tenant.agreement_pdf_path):
                    try:
                        os.remove(tenant.agreement_pdf_path)
                    except:
                        pass
                tenant.agreement_pdf_path = file_path
                logger.info(f"Set agreement_pdf_path to: {file_path}")
            elif file_type == "aadhar":
                # Delete old aadhar file if exists and different
                if tenant.aadhar_pdf_path and tenant.aadhar_pdf_path != file_path and os.path.exists(tenant.aadhar_pdf_path):
                    try:
                        os.remove(tenant.aadhar_pdf_path)
                    except:
                        pass
                tenant.aadhar_pdf_path = file_path
                logger.info(f"Set aadhar_pdf_path to: {file_path}")
            db.commit()
            db.refresh(tenant)
            logger.info(f"Tenant updated - agreement_pdf_path: {tenant.agreement_pdf_path}, aadhar_pdf_path: {tenant.aadhar_pdf_path}")
        else:
            logger.warning(f"Tenant {tenant_id} not found")

    return {
        "filename": file.filename,
        "file_path": file_path,
        "file_size": file_size,
        "file_type": file_type,
        "message": "File uploaded successfully"
    }


# IMPORTANT: This route MUST come BEFORE the /{file_path:path} route
# Otherwise FastAPI will match /tenant/2 as a file path
@router.get("/tenant/{tenant_id}")
def get_tenant_file(
    tenant_id: int,
    file_type: str = Query("agreement", description="Type of file: 'agreement' or 'aadhar'"),
    db: Session = Depends(get_db)
):
    """Get PDF file for a specific tenant"""
    logger.info(f"Get file request - tenant_id: {tenant_id}, file_type: '{file_type}'")

    tenant = db.query(Tenant).filter(Tenant.tenant_id == tenant_id).first()
    if not tenant:
        raise HTTPException(status_code=404, detail="Tenant not found")

    file_type = file_type.strip().lower()

    if file_type == "agreement":
        logger.info(f"Looking for agreement at: {tenant.agreement_pdf_path}")
        if not tenant.agreement_pdf_path:
            raise HTTPException(status_code=404, detail="Agreement PDF not found - no path stored")
        if not os.path.exists(tenant.agreement_pdf_path):
            raise HTTPException(status_code=404, detail=f"Agreement PDF file not found on disk: {tenant.agreement_pdf_path}")
        return FileResponse(tenant.agreement_pdf_path, media_type="application/pdf", filename=f"tenant_{tenant_id}_agreement.pdf")
    elif file_type == "aadhar":
        logger.info(f"Looking for aadhar at: {tenant.aadhar_pdf_path}")
        if not tenant.aadhar_pdf_path:
            raise HTTPException(status_code=404, detail="Aadhar PDF not found - no path stored")
        if not os.path.exists(tenant.aadhar_pdf_path):
            raise HTTPException(status_code=404, detail=f"Aadhar PDF file not found on disk: {tenant.aadhar_pdf_path}")
        return FileResponse(tenant.aadhar_pdf_path, media_type="application/pdf", filename=f"tenant_{tenant_id}_aadhar.pdf")
    else:
        raise HTTPException(status_code=400, detail="file_type must be 'agreement' or 'aadhar'")


# Agreement Generation Helper Functions
def number_to_words(num):
    """Convert number to words (Indian numbering system)"""
    ones = ['', 'One', 'Two', 'Three', 'Four', 'Five', 'Six', 'Seven', 'Eight', 'Nine',
            'Ten', 'Eleven', 'Twelve', 'Thirteen', 'Fourteen', 'Fifteen', 'Sixteen',
            'Seventeen', 'Eighteen', 'Nineteen']
    tens = ['', '', 'Twenty', 'Thirty', 'Forty', 'Fifty', 'Sixty', 'Seventy', 'Eighty', 'Ninety']

    if num == 0:
        return 'Zero'

    num = int(num)

    def two_digits(n):
        if n < 20:
            return ones[n]
        else:
            return tens[n // 10] + (' ' + ones[n % 10] if n % 10 else '')

    def three_digits(n):
        if n >= 100:
            return ones[n // 100] + ' Hundred' + (' and ' + two_digits(n % 100) if n % 100 else '')
        else:
            return two_digits(n)

    result = ''

    # Crores
    if num >= 10000000:
        result += three_digits(num // 10000000) + ' Crore '
        num %= 10000000

    # Lakhs
    if num >= 100000:
        result += two_digits(num // 100000) + ' Lakh '
        num %= 100000

    # Thousands
    if num >= 1000:
        result += two_digits(num // 1000) + ' Thousand '
        num %= 1000

    # Hundreds
    if num > 0:
        result += three_digits(num)

    return result.strip()


def date_to_words(d):
    """Convert date to format like '1st day of November 2025'"""
    # Ordinal suffixes
    def get_ordinal_suffix(day):
        if 11 <= day <= 13:
            return 'th'
        return {1: 'st', 2: 'nd', 3: 'rd'}.get(day % 10, 'th')

    months = ['', 'January', 'February', 'March', 'April', 'May', 'June',
              'July', 'August', 'September', 'October', 'November', 'December']

    day_ordinal = f"{d.day}{get_ordinal_suffix(d.day)}"
    month_word = months[d.month]

    return f"{day_ordinal} day of {month_word} {d.year}"


def calculate_duration(start_date, end_date):
    """Calculate duration between two dates"""
    delta = relativedelta(end_date, start_date)
    parts = []
    if delta.years > 0:
        parts.append(f"{delta.years} Year{'s' if delta.years > 1 else ''}")
    if delta.months > 0:
        parts.append(f"{delta.months} Month{'s' if delta.months > 1 else ''}")
    if delta.days > 0:
        parts.append(f"{delta.days} Day{'s' if delta.days > 1 else ''}")
    return ' '.join(parts) if parts else "0 Days"


def duration_to_words(start_date, end_date):
    """Convert duration to words"""
    delta = relativedelta(end_date, start_date)
    parts = []
    if delta.years > 0:
        parts.append(f"{number_to_words(delta.years)} Year{'s' if delta.years > 1 else ''}")
    if delta.months > 0:
        parts.append(f"{number_to_words(delta.months)} Month{'s' if delta.months > 1 else ''}")
    if delta.days > 0:
        parts.append(f"{number_to_words(delta.days)} Day{'s' if delta.days > 1 else ''}")
    return ' '.join(parts) if parts else "Zero Days"


def replace_placeholders_in_paragraph(paragraph, replacements):
    """Replace placeholders in a paragraph, handling split runs"""
    # First, try simple replacement within individual runs
    for placeholder, value in replacements.items():
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, str(value))

    # Now handle placeholders split across runs
    for placeholder, value in replacements.items():
        if placeholder in paragraph.text:
            # Placeholder exists in paragraph but wasn't replaced (split across runs)
            # We need to find and merge the runs containing the placeholder
            full_text = paragraph.text
            if placeholder not in full_text:
                continue

            # Find placeholder position in combined text
            start_idx = full_text.find(placeholder)
            if start_idx == -1:
                continue

            end_idx = start_idx + len(placeholder)

            # Find which runs contain the placeholder
            char_count = 0
            start_run_idx = None
            end_run_idx = None
            start_char_in_run = 0
            end_char_in_run = 0

            for i, run in enumerate(paragraph.runs):
                run_start = char_count
                run_end = char_count + len(run.text)

                if start_run_idx is None and run_start <= start_idx < run_end:
                    start_run_idx = i
                    start_char_in_run = start_idx - run_start

                if run_start < end_idx <= run_end:
                    end_run_idx = i
                    end_char_in_run = end_idx - run_start
                    break

                char_count = run_end

            if start_run_idx is not None and end_run_idx is not None:
                if start_run_idx == end_run_idx:
                    # Placeholder is in a single run
                    run = paragraph.runs[start_run_idx]
                    run.text = run.text[:start_char_in_run] + str(value) + run.text[end_char_in_run:]
                else:
                    # Placeholder spans multiple runs
                    # Put replacement text in first run, clear others
                    first_run = paragraph.runs[start_run_idx]
                    first_run.text = first_run.text[:start_char_in_run] + str(value)

                    # Clear middle runs
                    for i in range(start_run_idx + 1, end_run_idx):
                        paragraph.runs[i].text = ""

                    # Trim last run
                    last_run = paragraph.runs[end_run_idx]
                    last_run.text = last_run.text[end_char_in_run:]


def replace_placeholders_in_doc(doc, replacements):
    """Replace placeholders in entire document"""
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, replacements)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_placeholders_in_paragraph(paragraph, replacements)


@router.post("/generate-agreement/{tenant_id}")
def generate_agreement(
    tenant_id: int,
    owner_aadhar: Optional[str] = Query(None, description="Owner's Aadhar number (uses stored value if not provided)"),
    tenant_address: Optional[str] = Query(None, description="Tenant's permanent address (uses stored value if not provided)"),
    current_user: Client = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Generate rental agreement document from template for a specific tenant"""
    logger.info(f"Generate agreement request - tenant_id: {tenant_id}")

    # Fetch tenant with related data
    tenant = db.query(Tenant).filter(Tenant.tenant_id == tenant_id).first()
    if not tenant:
        raise HTTPException(status_code=404, detail="Tenant not found")

    # Fetch owner
    owner = db.query(Owner).filter(Owner.owner_id == tenant.owner_id).first()
    if not owner:
        raise HTTPException(status_code=404, detail="Owner not found")

    # Check access permission
    if current_user.role != "admin" and owner.client_id != current_user.client_id:
        raise HTTPException(status_code=403, detail="Access denied")

    # Fetch building
    building = db.query(Building).filter(Building.building_id == tenant.building_id).first()
    if not building:
        raise HTTPException(status_code=404, detail="Building not found")

    # Determine template based on building type
    template_type = "COMMERCIAL" if building.building_type == "Commercial" else "RESIDENTIAL"
    template_path = os.path.join(UPLOAD_DIR, "template", f"{template_type}.docx")

    if not os.path.exists(template_path):
        raise HTTPException(status_code=404, detail=f"Template not found: {template_path}")

    # Calculate rent values
    base_rent = float(tenant.rent_amount or 0)
    water_charge = float(tenant.water_charge or 0)
    maintenance_charge = float(tenant.maintenance_charge or 0)
    water_maintenance = water_charge + maintenance_charge
    total_rent = base_rent + water_maintenance
    advance = float(tenant.advance_amount or 0)

    # Use provided values or fall back to stored values
    final_owner_aadhar = owner_aadhar or owner.aadhar_number or "Not provided"
    final_tenant_address = tenant_address or tenant.address or "Address not provided"

    # Prepare replacements
    today = date.today()
    replacements = {
        # Agreement start date in words (e.g., "1st day of November 2025")
        "{AGREEMENT_START_DATE_IN_WORDS}": date_to_words(tenant.agreement_start_date),

        # Owner details
        "{OWNER_NAME}": owner.name,
        "{OWNER_AADHAR}": final_owner_aadhar,
        "{OWNER_ADDRESS}": owner.address or building.location or "Address not provided",

        # Tenant details
        "{TENANT_NAME}": tenant.name,
        "{TENANT_AADHAR}": tenant.aadhar_number or "Not provided",
        "{TENANT_ADDRESS}": final_tenant_address,

        # Rent details
        "{BASE_RENT}": f"{base_rent:,.0f}",
        "{BASE_RENT_IN_WORDS}": number_to_words(base_rent),
        "{WATER_MAINTENANCE}": "Water and Maintenance Charges" if water_maintenance > 0 else "Additional Charges",
        "{WATER_MAINTENANCE_RENT}": f"{water_maintenance:,.0f}",
        "{TOTAL_RENT}": f"{total_rent:,.0f}",
        "{TOTAL_RENT_IN_WORDS}": number_to_words(total_rent),
        "{RENT_DUE_DATE}": str(tenant.rent_due_date),

        # Advance
        "{ADVANCE}": f"{advance:,.0f}",
        "{ADVANCE_IN_WORDS}": number_to_words(advance),

        # Agreement dates
        "{AGREEMENT_START_DATE}": tenant.agreement_start_date.strftime("%d-%m-%Y"),
        "{AGREEMENT_END_DATE}": tenant.agreement_end_date.strftime("%d-%m-%Y"),
        "{AGREEMENT_DURATION}": calculate_duration(tenant.agreement_start_date, tenant.agreement_end_date),
        "{AGREEMENT_DURATION_IN_WORDS}": duration_to_words(tenant.agreement_start_date, tenant.agreement_end_date),
    }

    # Load and process template
    try:
        doc = Document(template_path)
        replace_placeholders_in_doc(doc, replacements)

        # Save generated document
        generated_dir = os.path.join(UPLOAD_DIR, "generated")
        Path(generated_dir).mkdir(parents=True, exist_ok=True)

        output_filename = f"agreement_{tenant_id}_{tenant.name.replace(' ', '_')}_{today.strftime('%Y%m%d')}.docx"
        output_path = os.path.join(generated_dir, output_filename)

        doc.save(output_path)
        logger.info(f"Agreement generated: {output_path}")

        return FileResponse(
            output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=output_filename
        )

    except Exception as e:
        logger.error(f"Error generating agreement: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error generating agreement: {str(e)}")


@router.get("/agreement-preview/{tenant_id}")
def get_agreement_preview_data(
    tenant_id: int,
    current_user: Client = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get data for agreement preview before generation"""
    tenant = db.query(Tenant).filter(Tenant.tenant_id == tenant_id).first()
    if not tenant:
        raise HTTPException(status_code=404, detail="Tenant not found")

    owner = db.query(Owner).filter(Owner.owner_id == tenant.owner_id).first()
    if not owner:
        raise HTTPException(status_code=404, detail="Owner not found")

    # Check access permission
    if current_user.role != "admin" and owner.client_id != current_user.client_id:
        raise HTTPException(status_code=403, detail="Access denied")

    building = db.query(Building).filter(Building.building_id == tenant.building_id).first()

    base_rent = float(tenant.rent_amount or 0)
    water_charge = float(tenant.water_charge or 0)
    maintenance_charge = float(tenant.maintenance_charge or 0)
    total_rent = base_rent + water_charge + maintenance_charge

    return {
        "tenant": {
            "tenant_id": tenant.tenant_id,
            "name": tenant.name,
            "phone": tenant.phone,
            "email": tenant.email,
            "aadhar_number": tenant.aadhar_number,
            "address": tenant.address,
            "portion_number": tenant.portion_number,
        },
        "owner": {
            "owner_id": owner.owner_id,
            "name": owner.name,
            "phone": owner.phone,
            "email": owner.email,
            "address": owner.address,
            "aadhar_number": owner.aadhar_number,
        },
        "building": {
            "building_id": building.building_id if building else None,
            "name": building.building_name if building else None,
            "type": building.building_type if building else None,
            "location": building.location if building else None,
        },
        "rent": {
            "base_rent": base_rent,
            "water_charge": water_charge,
            "maintenance_charge": maintenance_charge,
            "total_rent": total_rent,
            "advance_amount": float(tenant.advance_amount or 0),
            "rent_due_date": tenant.rent_due_date,
        },
        "agreement": {
            "start_date": tenant.agreement_start_date.isoformat(),
            "end_date": tenant.agreement_end_date.isoformat(),
            "duration": calculate_duration(tenant.agreement_start_date, tenant.agreement_end_date),
        }
    }


# This catch-all route MUST come LAST
@router.get("/download/{file_path:path}")
def get_file(file_path: str):
    """Get a file by path"""
    # Security: Ensure file is within upload directory
    full_path = os.path.abspath(os.path.join(UPLOAD_DIR, file_path))
    upload_dir_abs = os.path.abspath(UPLOAD_DIR)

    if not full_path.startswith(upload_dir_abs):
        raise HTTPException(status_code=403, detail="Access denied")

    if not os.path.exists(full_path):
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(full_path, media_type="application/pdf")
